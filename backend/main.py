from __future__ import annotations

import io
import json
import logging
import logging.handlers
import os
import base64
import shutil
import subprocess
import sys
import tempfile
import threading
import time
import uuid
import zipfile
from dataclasses import dataclass, field
from errno import ENOSPC
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

import anyio
from fastapi import BackgroundTasks, Depends, FastAPI, File, Form, Header, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from starlette.exceptions import HTTPException as StarletteHTTPException
from starlette.responses import JSONResponse

from app.auth.deps import CurrentAuthorizedUser
from app.auth.router import router as auth_router

try:
    from dotenv import load_dotenv

    _dotenv_path = Path(__file__).resolve().parent / ".env"
    if _dotenv_path.exists():
        load_dotenv(dotenv_path=_dotenv_path)
    else:
        load_dotenv()
except Exception:
    pass


LOG_DIR = Path(os.getenv("CLIPBUILDER_LOG_DIR", Path(__file__).resolve().parent / "logs"))
LOG_DIR.mkdir(parents=True, exist_ok=True)

logger = logging.getLogger("clipbuilder")
if not logger.handlers:
    logger.setLevel(logging.INFO)
    fmt = logging.Formatter(
        fmt="%(asctime)s %(levelname)s %(name)s %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )

    info_handler = logging.handlers.RotatingFileHandler(
        LOG_DIR / "clipbuilder.log",
        maxBytes=2_000_000,
        backupCount=5,
        encoding="utf-8",
    )
    info_handler.setLevel(logging.INFO)
    info_handler.setFormatter(fmt)

    error_handler = logging.handlers.RotatingFileHandler(
        LOG_DIR / "clipbuilder.error.log",
        maxBytes=2_000_000,
        backupCount=10,
        encoding="utf-8",
    )
    error_handler.setLevel(logging.ERROR)
    error_handler.setFormatter(fmt)

    logger.addHandler(info_handler)
    logger.addHandler(error_handler)


def _http_exception_handler(_request: Request, exc: StarletteHTTPException):
    # Starlette hides multipart errors behind a generic message; return actionable hint.
    if exc.status_code == 400 and str(exc.detail) == "There was an error parsing the body":
        logger.warning("multipart parse error: %s", exc.detail)
        return JSONResponse(
            status_code=400,
            content={
                "detail": (
                    "Falha ao ler o upload (multipart). Isso costuma acontecer com vídeos grandes por limite do parser. "
                    "Use o upload raw (endpoint /videos/raw) no frontend/backend atualizados."
                )
            },
        )

    if exc.status_code >= 500:
        logger.error("http error %s: %s", exc.status_code, exc.detail, exc_info=True)
    elif exc.status_code >= 400:
        logger.warning("http %s: %s", exc.status_code, exc.detail)
    return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})


def _unhandled_exception_handler(_request: Request, exc: Exception):
    logger.error("unhandled error: %s", exc, exc_info=True)
    return JSONResponse(status_code=500, content={"detail": "Erro interno no servidor"})


MAX_TOTAL_IMAGE_BYTES = 50 * 1024 * 1024  # 50MB
MAX_SINGLE_IMAGE_BYTES = 15 * 1024 * 1024  # 15MB


def _env_int(name: str, default: int) -> int:
    raw = (os.getenv(name) or "").strip()
    if not raw:
        return default
    try:
        value = int(raw)
    except ValueError:
        return default
    return value if value > 0 else default


MAX_VIDEO_BYTES = _env_int("CLIPBUILDER_MAX_VIDEO_BYTES", 700 * 1024 * 1024)  # default 700MB
MAX_VIDEO_MB = max(1, int(MAX_VIDEO_BYTES / (1024 * 1024)))

logger.info("config: CLIPBUILDER_MAX_VIDEO_BYTES=%s (~%s MB)", MAX_VIDEO_BYTES, MAX_VIDEO_MB)

GEMINI_CLIP_SECONDS = _env_int("CLIPBUILDER_GEMINI_CLIP_SECONDS", 90)
GEMINI_PROXY_HEIGHT = _env_int("CLIPBUILDER_GEMINI_PROXY_HEIGHT", 720)

DEFAULT_GEMINI_MODEL = os.getenv("CLIPBUILDER_GEMINI_MODEL", "models/gemini-2.5-flash")
GEMINI_POLL_TIMEOUT_SECONDS = 300
GEMINI_POLL_INTERVAL_SECONDS = 2

AI_LANGUAGE = (os.getenv("CLIPBUILDER_AI_LANGUAGE") or "pt-BR").strip() or "pt-BR"

# Groq API configuration (Llama 4 Vision + Whisper Turbo)
GROQ_API_KEY = (os.getenv("GROQ_API_KEY") or "").strip()
DEFAULT_GROQ_VISION_MODEL = os.getenv("CLIPBUILDER_GROQ_VISION_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct")
DEFAULT_GROQ_WHISPER_MODEL = os.getenv("CLIPBUILDER_GROQ_WHISPER_MODEL", "whisper-large-v3-turbo")
GROQ_FRAME_COUNT = _env_int("CLIPBUILDER_GROQ_FRAME_COUNT", 5)  # Max 5 images per Groq request

YTDLP_COOKIES_FILE = (os.getenv("CLIPBUILDER_YTDLP_COOKIES_FILE") or "").strip()
YTDLP_COOKIES_FROM_BROWSER = (os.getenv("CLIPBUILDER_YTDLP_COOKIES_FROM_BROWSER") or "").strip()
YTDLP_COOKIES_FROM_BROWSER_ARGS = (os.getenv("CLIPBUILDER_YTDLP_COOKIES_FROM_BROWSER_ARGS") or "").strip()

app = FastAPI(title="ClipBuilder")
app.include_router(auth_router, prefix="/auth")

app.add_exception_handler(StarletteHTTPException, _http_exception_handler)
app.add_exception_handler(Exception, _unhandled_exception_handler)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@dataclass
class VideoEntry:
    path: Path
    status: str  # ready|error
    clip_cache: dict[str, str] = field(default_factory=dict)  # key -> gemini_file_name
    error: str | None = None


DATA_DIR = Path(os.getenv("CLIPBUILDER_DATA_DIR", Path(__file__).resolve().parent / "data"))
DATA_DIR.mkdir(parents=True, exist_ok=True)
logger.info("config: CLIPBUILDER_DATA_DIR=%s", DATA_DIR)

# Maximum number of video files to keep in DATA_DIR (oldest files are removed when exceeded)
MAX_DATA_FILES = _env_int("CLIPBUILDER_MAX_DATA_FILES", 5)
logger.info("config: CLIPBUILDER_MAX_DATA_FILES=%s", MAX_DATA_FILES)

_videos_lock = threading.Lock()
_videos: dict[str, VideoEntry] = {}


def _cleanup_old_files() -> None:
    """Remove oldest video files from DATA_DIR if count exceeds MAX_DATA_FILES.
    
    Files are sorted by modification time and the oldest ones are deleted
    until only MAX_DATA_FILES remain.
    """
    try:
        # Get all video files in DATA_DIR
        video_files = [
            f for f in DATA_DIR.iterdir()
            if f.is_file() and f.name.startswith("video_") and f.suffix.lower() in {".mp4", ".mkv"}
        ]
        
        if len(video_files) <= MAX_DATA_FILES:
            return
        
        # Sort by modification time (oldest first)
        video_files.sort(key=lambda f: f.stat().st_mtime)
        
        # Remove oldest files until we're at the limit
        files_to_remove = len(video_files) - MAX_DATA_FILES
        for i in range(files_to_remove):
            file_to_delete = video_files[i]
            try:
                # Also remove from in-memory cache
                video_id = file_to_delete.stem.replace("video_", "")
                with _videos_lock:
                    if video_id in _videos:
                        del _videos[video_id]
                
                file_to_delete.unlink()
                logger.info("cleanup: removed old video file %s", file_to_delete.name)
            except Exception as exc:
                logger.warning("cleanup: failed to remove %s: %s", file_to_delete.name, exc)
    except Exception as exc:
        logger.warning("cleanup: failed to clean old files: %s", exc)


def _restore_video_entry_if_missing(video_id: str) -> VideoEntry | None:
    """Restore an in-memory entry from disk after a server reload.

    We intentionally keep _videos in-memory for simplicity, but uvicorn --reload
    restarts the process, so we reconstruct entries on demand when possible.
    """
    candidates = [DATA_DIR / f"video_{video_id}.mp4", DATA_DIR / f"video_{video_id}.mkv"]
    for path in candidates:
        try:
            if path.exists() and path.is_file() and path.stat().st_size > 0:
                entry = VideoEntry(path=path, status="ready")
                _videos[video_id] = entry
                return entry
        except Exception:
            continue
    return None

def _configure_genai(api_key: str) -> None:
    try:
        import google.generativeai as genai
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="Dependência 'google-generativeai' não instalada no backend.",
        ) from exc

    genai.configure(api_key=api_key)


def _get_api_key(x_google_api_key: str | None) -> str:
    api_key = (os.getenv("GOOGLE_API_KEY") or "").strip() or (x_google_api_key or "").strip()
    if not api_key:
        raise HTTPException(
            status_code=400,
            detail="GOOGLE_API_KEY não configurada. Defina no backend (.env) ou envie no header X-Google-Api-Key.",
        )
    return api_key


def _upload_video_to_gemini(video_path: Path, api_key: str) -> str:
    _configure_genai(api_key)
    import google.generativeai as genai

    uploaded = genai.upload_file(path=str(video_path))
    file_name = getattr(uploaded, "name", None)
    if not file_name:
        raise RuntimeError("Falha ao fazer upload do vídeo para o Gemini")

    deadline = time.time() + GEMINI_POLL_TIMEOUT_SECONDS
    while time.time() < deadline:
        current = genai.get_file(file_name)
        state = getattr(getattr(current, "state", None), "name", None) or str(getattr(current, "state", ""))
        if state == "ACTIVE":
            return file_name
        if state in {"FAILED", "ERROR"}:
            raise RuntimeError("Gemini falhou ao processar o arquivo de vídeo")
        time.sleep(GEMINI_POLL_INTERVAL_SECONDS)

    raise RuntimeError("Timeout aguardando o Gemini processar o arquivo (ACTIVE)")


def _format_timestamp(seconds: float) -> str:
    if seconds < 0:
        seconds = 0
    total = int(seconds)
    hh = total // 3600
    mm = (total % 3600) // 60
    ss = total % 60
    return f"{hh:02d}:{mm:02d}:{ss:02d}"


def _parse_timestamp_to_seconds(timestamp: str) -> float:
    ts = (timestamp or "").strip()
    if not ts:
        return 0.0
    parts = ts.split(":")
    if len(parts) > 3:
        raise ValueError("timestamp inválido")
    try:
        nums = [float(p) for p in parts]
    except ValueError as exc:
        raise ValueError("timestamp inválido") from exc

    while len(nums) < 3:
        nums.insert(0, 0.0)
    hh, mm, ss = nums
    return max(0.0, hh * 3600.0 + mm * 60.0 + ss)


def _ensure_ffmpeg() -> str:
    path = shutil.which("ffmpeg")
    if not path:
        raise HTTPException(
            status_code=500,
            detail="ffmpeg não está instalado no servidor. No Fedora: sudo dnf install -y ffmpeg",
        )
    return path


def _is_supported_youtube_url(url: str) -> bool:
    try:
        parsed = urlparse(url)
    except Exception:
        return False
    if parsed.scheme not in {"http", "https"}:
        return False
    host = (parsed.netloc or "").lower()
    if host.startswith("www."):
        host = host[4:]
    return host in {"youtube.com", "m.youtube.com", "youtu.be"} or host.endswith(".youtube.com")


def _ytdlp_cmd() -> list[str]:
    # Prefer binary if available, otherwise fall back to python module.
    if shutil.which("yt-dlp"):
        return ["yt-dlp"]
    return [sys.executable, "-m", "yt_dlp"]


def _ytdlp_cookie_args() -> list[str]:
    raw = (YTDLP_COOKIES_FILE or "").strip()
    if not raw:
        return []
    try:
        path = Path(raw).expanduser()
        if not path.is_absolute():
            path = (Path(__file__).resolve().parent / path).resolve()
        if path.exists() and path.is_file():
            return ["--cookies", str(path)]
    except Exception:
        pass
    return []


def _ytdlp_browser_cookie_args() -> list[str]:
    raw = (YTDLP_COOKIES_FROM_BROWSER or "").strip()
    if not raw:
        return []
    # Allow passing extra selector arguments if needed (e.g. "firefox:default" or keyring selector).
    if (YTDLP_COOKIES_FROM_BROWSER_ARGS or "").strip():
        selector = f"{raw}:{YTDLP_COOKIES_FROM_BROWSER_ARGS.strip()}"
    else:
        selector = raw
    return ["--cookies-from-browser", selector]


def _ytdlp_auth_args() -> list[str]:
    # Prefer browser cookies when configured (avoids manual export), fall back to cookies file.
    args = _ytdlp_browser_cookie_args()
    if args:
        return args
    return _ytdlp_cookie_args()


def _download_youtube_video(*, url: str, out_path: Path) -> None:
    _ensure_ffmpeg()  # yt-dlp may need it to merge streams

    cmd = _ytdlp_cmd() + [
        "--no-playlist",
        "--no-warnings",
        *(_ytdlp_auth_args()),
        "--max-filesize",
        f"{MAX_VIDEO_MB}M",
        "-f",
        # Prefer MP4; fall back to best.
        "bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best",
        "--merge-output-format",
        "mp4",
        "-o",
        str(out_path),
        url,
    ]

    proc = subprocess.run(cmd, capture_output=True, text=True)
    if proc.returncode != 0:
        stderr = (proc.stderr or "").strip()
        stdout = (proc.stdout or "").strip()
        tail = "\n".join((stderr.splitlines() + stdout.splitlines())[-20:]).strip()

        lowered = (tail or "").lower()
        if "sign in to confirm" in lowered and "not a bot" in lowered:
            hint = (
                "O YouTube bloqueou o download e pediu verificação (\"not a bot\"). "
                "Para vídeos assim, é necessário fornecer cookies do seu navegador para o yt-dlp. "
                "Opção A: exporte um cookies.txt e configure CLIPBUILDER_YTDLP_COOKIES_FILE. "
                "Opção B: configure CLIPBUILDER_YTDLP_COOKIES_FROM_BROWSER (ex.: firefox, chrome, chromium) para usar o perfil do navegador."
            )
            raise HTTPException(status_code=403, detail=hint)

        raise RuntimeError(f"yt-dlp failed ({proc.returncode}): {tail or 'unknown error'}")

    if not out_path.exists() or out_path.stat().st_size == 0:
        raise RuntimeError("Download falhou: arquivo de saída não foi criado")

    if out_path.stat().st_size > MAX_VIDEO_BYTES:
        try:
            out_path.unlink(missing_ok=True)
        except Exception:
            pass
        raise HTTPException(
            status_code=413,
            detail=(
                f"Vídeo muito grande (limite atual: {MAX_VIDEO_MB} MB). "
                "Ajuste CLIPBUILDER_MAX_VIDEO_BYTES no backend/.env."
            ),
        )


def _make_gemini_clip(*, source_path: Path, timestamp_seconds: float, clip_seconds: int, out_path: Path) -> tuple[int, int]:
    ffmpeg: str = _ensure_ffmpeg()
    clip_seconds = max(10, int(clip_seconds))
    half = clip_seconds // 2
    start = max(0, int(timestamp_seconds) - half)
    duration = clip_seconds

    def run(cmd: list[str]) -> subprocess.CompletedProcess[str]:
        return subprocess.run(cmd, capture_output=True, text=True)

    def tail(stderr: str) -> str:
        return "\n".join((stderr or "").strip().splitlines()[-20:]).strip()

    # Attempt 1: re-encode to a smaller proxy (best for cost/speed), but may fail on
    # ffmpeg builds without libx264 (common on Fedora ffmpeg-free).
    vf = f"scale=-2:{max(144, int(GEMINI_PROXY_HEIGHT))}"
    for video_encoder in ["libx264", "libopenh264"]:
        cmd = [
            ffmpeg,
            "-y",
            "-ss",
            str(start),
            "-t",
            str(duration),
            "-i",
            str(source_path),
            "-vf",
            vf,
            "-c:v",
            video_encoder,
            "-preset",
            "veryfast",
            "-crf",
            "28",
            "-c:a",
            "aac",
            "-b:a",
            "96k",
            "-movflags",
            "+faststart",
            str(out_path),
        ]

        proc = run(cmd)
        if proc.returncode == 0:
            return start, duration

        err = proc.stderr or ""
        # If encoder isn't available, try the next option.
        if f"Unknown encoder '{video_encoder}'" in err:
            continue

        # Other failures: break and try stream copy.
        break

    # Attempt 2: stream copy (no re-encode). This is very compatible and avoids encoder issues.
    copy_cmd = [
        ffmpeg,
        "-y",
        "-ss",
        str(start),
        "-t",
        str(duration),
        "-i",
        str(source_path),
        "-c",
        "copy",
        "-movflags",
        "+faststart",
        str(out_path),
    ]

    proc = run(copy_cmd)
    if proc.returncode != 0:
        raise RuntimeError("ffmpeg falhou: " + tail(proc.stderr or ""))
    return start, duration


def _describe_at_timestamp(*, gemini_file_name: str, timestamp: str, clip_seconds: int, api_key: str, model_name: str, user_prompt: str | None = None, include_timestamp: bool = True) -> str:
    _configure_genai(api_key)
    import google.generativeai as genai
    from google.api_core import exceptions as gexc  # type: ignore

    # Accept either "gemini-2.0-flash" or "models/gemini-2.0-flash".
    normalized_model = (model_name or "").strip()
    if normalized_model and not normalized_model.startswith("models/"):
        normalized_model = f"models/{normalized_model}"

    file_ref = genai.get_file(gemini_file_name)

    # 1. Instrução de Sistema / Persona
    system_instruction = (
        f"Você é um especialista em Documentação Técnica de Software. "
        f"Seu objetivo é criar tutoriais passo a passo claros e diretos. "
        f"Responda sempre em português do Brasil (pt-BR). Idioma preferido: {AI_LANGUAGE}."
    )

    # 2. Definição de Contexto Temporal (análise expandida)
    time_instruction = (
        f"O foco principal é o timestamp: {timestamp}. "
        f"No entanto, para entender o contexto, ANALISE o vídeo desde 40 segundos ANTES até 40 segundos DEPOIS desse momento. "
        "Use esse intervalo para identificar qual processo lógico está sendo iniciado ou concluído. "
        "Entenda o fluxo: Onde o usuário clicou antes para chegar aqui? O que acontece depois que confirma a ação? "
        "Identifique o OBJETIVO FINAL da ação (ex: não é apenas 'preencher campo', é 'Configurar filtro de data')."
    )

    # 3. Regras de Formatação e Estilo
    formatting_instruction = (
        "REGRAS DE SAÍDA:\n"
        "1. Identifique a ação macro (ex: 'Cadastrando um novo usuário').\n"
        "2. Gere APENAS os passos imperativos necessários para realizar essa ação.\n"
        "3. Ignore movimentos de mouse erráticos ou tentativas falhas.\n"
        "4. Não use narração (ex: 'O usuário clica...'). Use imperativo (ex: 'Clique em Salvar').\n"
        "5. Seja conciso (3 a 10 passos).\n"
        "6. Evite detalhes técnicos visuais (coordenadas X/Y, posições específicas como 'célula A1') a menos que cruciais.\n"
        "7. Não mencione 'vídeo', 'clipe', 'cena', 'tela', 'o usuário' ou 'o mouse'.\n"
        "8. Foque nos comandos, menus e opções relevantes para reproduzir o processo."
    )

    base_parts: list[str] = [
        system_instruction,
        time_instruction,
        formatting_instruction,
    ]

    if include_timestamp:
        base_parts.append(f"O procedimento ocorre especificamente ao redor de {timestamp}.")
    else:
        base_parts.append("Descreva o procedimento exibido. Não inclua timestamp na resposta.")

    base = "\n\n".join(base_parts)

    if user_prompt and str(user_prompt).strip():
        prompt = base + "\n\n" + "Contexto extra do usuário (se aplicável):\n" + str(user_prompt).strip()
    else:
        prompt = base

    model = genai.GenerativeModel(normalized_model or DEFAULT_GEMINI_MODEL)

    def _extract_retry_seconds(message: str) -> float | None:
        # Example: "Please retry in 13.644857575s."
        import re

        m = re.search(r"retry in\s+([0-9]+(?:\.[0-9]+)?)s", message, flags=re.IGNORECASE)
        if not m:
            return None
        try:
            value = float(m.group(1))
        except ValueError:
            return None
        return value if value > 0 else None

    response = None
    for attempt in range(2):
        try:
            response = model.generate_content([file_ref, prompt])
            break
        except (getattr(gexc, "ResourceExhausted", Exception), getattr(gexc, "TooManyRequests", Exception)) as exc:
            if attempt >= 1:
                raise
            retry_seconds = _extract_retry_seconds(str(exc) or "")
            # Cap to keep requests responsive.
            sleep_for = min(max(retry_seconds or 2.0, 0.5), 15.0)
            import time

            time.sleep(sleep_for)

    if response is None:
        return ""
    text = getattr(response, "text", None)
    if not text:
        return ""
    return str(text).strip()


# ---------------------------------------------------------------------------
# Groq API Integration (Llama 4 Vision + Whisper Turbo)
# ---------------------------------------------------------------------------

def _get_groq_api_key(header_key: str | None = None) -> str:
    """Get Groq API key from environment or header."""
    api_key = GROQ_API_KEY or (header_key or "").strip()
    if not api_key:
        raise HTTPException(
            status_code=400,
            detail="GROQ_API_KEY não configurada. Defina no backend (.env) ou envie no header X-Groq-Api-Key.",
        )
    return api_key


def _extract_frames_from_video(
    video_path: Path,
    timestamp_seconds: float,
    window_seconds: int = 40,
    frame_count: int = 5,
) -> list[bytes]:
    """Extract frames from video around the timestamp.
    
    Extracts `frame_count` frames evenly distributed in a window of
    `window_seconds` before and after the timestamp.
    Returns list of PNG image bytes.
    """
    ffprobe_cmd = [
        "ffprobe",
        "-v", "error",
        "-show_entries", "format=duration",
        "-of", "default=noprint_wrappers=1:nokey=1",
        str(video_path),
    ]
    try:
        result = subprocess.run(ffprobe_cmd, capture_output=True, text=True, check=True)
        video_duration = float(result.stdout.strip())
    except Exception:
        video_duration = timestamp_seconds + window_seconds + 10

    # Calculate frame timestamps
    start_time = max(0, timestamp_seconds - window_seconds)
    end_time = min(video_duration, timestamp_seconds + window_seconds)
    actual_window = end_time - start_time
    
    if frame_count <= 1:
        timestamps = [timestamp_seconds]
    else:
        step = actual_window / (frame_count - 1)
        timestamps = [start_time + i * step for i in range(frame_count)]
    
    frames: list[bytes] = []
    for ts in timestamps:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            ffmpeg_cmd = [
                "ffmpeg",
                "-y",
                "-ss", str(ts),
                "-i", str(video_path),
                "-vframes", "1",
                "-vf", f"scale=-1:{GEMINI_PROXY_HEIGHT}",  # Reuse existing height config
                "-f", "image2",
                tmp_path,
            ]
            subprocess.run(ffmpeg_cmd, capture_output=True, check=True)
            
            with open(tmp_path, "rb") as f:
                frame_data = f.read()
                if frame_data:
                    frames.append(frame_data)
        except Exception as exc:
            logger.warning("Failed to extract frame at %.2fs: %s", ts, exc)
        finally:
            try:
                Path(tmp_path).unlink(missing_ok=True)
            except Exception:
                pass
    
    return frames


def _extract_frames_from_gif(gif_bytes: bytes, frame_count: int = 5) -> list[bytes]:
    """Extract frames from a GIF file using Pillow.
    
    Returns up to `frame_count` frames evenly distributed across the GIF,
    as PNG bytes for sending to Groq Vision.
    """
    try:
        from PIL import Image
        from PIL import ImageSequence
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="Dependência 'Pillow' não instalada. Execute: pip install Pillow",
        ) from exc

    try:
        img = Image.open(io.BytesIO(gif_bytes))
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Arquivo GIF inválido ou corrompido: {exc!s}",
        ) from exc

    if getattr(img, "format", None) and str(img.format).upper() not in ("GIF",):
        raise HTTPException(
            status_code=400,
            detail="O arquivo não é um GIF válido.",
        )

    # Collect all frames (ImageSequence.Iterator consumes, so we list once)
    seq_frames: list[Any] = []
    for frame in ImageSequence.Iterator(img):
        seq_frames.append(frame.copy())
    if not seq_frames:
        seq_frames = [img.copy()]

    n = len(seq_frames)
    if frame_count <= 1 or n <= 1:
        indices = [0]
    else:
        step = (n - 1) / max(1, frame_count - 1)
        indices = [min(int(round(i * step)), n - 1) for i in range(frame_count)]
        indices = sorted(set(indices))[:frame_count]

    out: list[bytes] = []
    for idx in indices:
        try:
            frame_img = seq_frames[idx]
            buf = io.BytesIO()
            frame_img.save(buf, format="PNG")
            data = buf.getvalue()
            if data:
                out.append(data)
        except Exception as exc:
            logger.warning("Failed to export GIF frame %s as PNG: %s", idx, exc)

    if not out:
        raise HTTPException(
            status_code=400,
            detail="Não foi possível extrair frames do GIF.",
        )
    return out


def _extract_audio_clip(
    video_path: Path,
    timestamp_seconds: float,
    window_seconds: int = 40,
) -> Path | None:
    """Extract audio clip from video around the timestamp.
    
    Returns path to temporary WAV file, or None if extraction fails.
    """
    # Get video duration
    ffprobe_cmd = [
        "ffprobe",
        "-v", "error",
        "-show_entries", "format=duration",
        "-of", "default=noprint_wrappers=1:nokey=1",
        str(video_path),
    ]
    try:
        result = subprocess.run(ffprobe_cmd, capture_output=True, text=True, check=True)
        video_duration = float(result.stdout.strip())
    except Exception:
        video_duration = timestamp_seconds + window_seconds + 10
    
    start_time = max(0, timestamp_seconds - window_seconds)
    end_time = min(video_duration, timestamp_seconds + window_seconds)
    duration = end_time - start_time
    
    # Create temp file for audio
    audio_path = Path(tempfile.mktemp(suffix=".wav"))
    
    ffmpeg_cmd = [
        "ffmpeg",
        "-y",
        "-ss", str(start_time),
        "-i", str(video_path),
        "-t", str(duration),
        "-vn",  # No video
        "-acodec", "pcm_s16le",
        "-ar", "16000",  # 16kHz for Whisper
        "-ac", "1",  # Mono
        str(audio_path),
    ]
    
    try:
        subprocess.run(ffmpeg_cmd, capture_output=True, check=True)
        if audio_path.exists() and audio_path.stat().st_size > 1000:
            return audio_path
    except Exception as exc:
        logger.warning("Failed to extract audio: %s", exc)
    
    try:
        audio_path.unlink(missing_ok=True)
    except Exception:
        pass
    return None


def _transcribe_with_groq(audio_path: Path, api_key: str) -> str:
    """Transcribe audio using Groq Whisper API."""
    try:
        from groq import Groq
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="Dependência 'groq' não instalada. Execute: pip install groq",
        ) from exc
    
    client = Groq(api_key=api_key)
    
    try:
        with open(audio_path, "rb") as audio_file:
            transcription = client.audio.transcriptions.create(
                file=(audio_path.name, audio_file.read()),
                model=DEFAULT_GROQ_WHISPER_MODEL,
                language="pt",  # Portuguese
            )
        return transcription.text or ""
    except Exception as exc:
        logger.warning("Groq transcription failed: %s", exc)
        return ""


def _analyze_frames_with_groq(
    frames: list[bytes],
    prompt: str,
    api_key: str,
    model: str | None = None,
) -> str:
    """Analyze frames using Groq Vision API (Llama 4)."""
    try:
        from groq import Groq
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="Dependência 'groq' não instalada. Execute: pip install groq",
        ) from exc
    
    if not frames:
        raise HTTPException(status_code=400, detail="Nenhum frame extraído do vídeo")
    
    client = Groq(api_key=api_key)
    vision_model = model or DEFAULT_GROQ_VISION_MODEL
    
    # Build message content with images
    content: list[dict[str, Any]] = [{"type": "text", "text": prompt}]
    
    for i, frame_bytes in enumerate(frames[:5]):  # Max 5 images per Groq request
        b64_image = base64.b64encode(frame_bytes).decode("ascii")
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/png;base64,{b64_image}",
            },
        })
    
    try:
        response = client.chat.completions.create(
            model=vision_model,
            messages=[{"role": "user", "content": content}],
            temperature=0.7,
            max_completion_tokens=2048,
        )
        return response.choices[0].message.content or ""
    except Exception as exc:
        error_msg = str(exc).lower()
        if "rate" in error_msg or "limit" in error_msg or "429" in error_msg:
            raise HTTPException(
                status_code=429,
                detail="Limite de requisições do Groq excedido. Tente novamente em alguns segundos.",
            ) from exc
        if "invalid" in error_msg and "api" in error_msg:
            raise HTTPException(
                status_code=401,
                detail="Chave da API Groq inválida. Verifique GROQ_API_KEY no .env",
            ) from exc
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao analisar com Groq: {str(exc)[:200]}",
        ) from exc


def _describe_with_groq(
    *,
    video_path: Path,
    timestamp: str,
    timestamp_seconds: float,
    api_key: str,
    model_name: str | None = None,
    user_prompt: str | None = None,
    include_timestamp: bool = True,
) -> str:
    """Describe video content using Groq Vision + Whisper.
    
    Extracts frames and audio from the video, transcribes audio,
    and uses vision model to generate description combining both contexts.
    """
    # 1. Extract frames
    frames = _extract_frames_from_video(
        video_path=video_path,
        timestamp_seconds=timestamp_seconds,
        window_seconds=40,
        frame_count=GROQ_FRAME_COUNT,
    )
    
    if not frames:
        raise HTTPException(status_code=500, detail="Não foi possível extrair frames do vídeo")
    
    # 2. Extract and transcribe audio (in parallel would be better, but keeping simple)
    audio_context = ""
    audio_path = _extract_audio_clip(video_path, timestamp_seconds, window_seconds=40)
    if audio_path:
        try:
            audio_context = _transcribe_with_groq(audio_path, api_key)
        finally:
            try:
                audio_path.unlink(missing_ok=True)
            except Exception:
                pass
    
    # 3. Build prompt (same style as Gemini for consistency)
    system_instruction = (
        f"Você é um especialista em Documentação Técnica de Software. "
        f"Seu objetivo é criar tutoriais passo a passo claros e diretos. "
        f"Responda sempre em português do Brasil (pt-BR). Idioma preferido: {AI_LANGUAGE}."
    )
    
    time_instruction = (
        f"O foco principal é o timestamp: {timestamp}. "
        f"Você está vendo {len(frames)} frames extraídos do vídeo ao redor desse momento. "
        "Analise a sequência de imagens para entender o fluxo: O que está sendo feito? Qual o objetivo final?"
    )
    
    formatting_instruction = (
        "REGRAS DE SAÍDA:\n"
        "1. Identifique a ação macro (ex: 'Cadastrando um novo usuário').\n"
        "2. Gere APENAS os passos imperativos necessários para realizar essa ação.\n"
        "3. Ignore movimentos de mouse erráticos ou tentativas falhas.\n"
        "4. Não use narração (ex: 'O usuário clica...'). Use imperativo (ex: 'Clique em Salvar').\n"
        "5. Seja conciso (3 a 10 passos).\n"
        "6. Evite detalhes técnicos visuais (coordenadas X/Y, posições específicas).\n"
        "7. Não mencione 'vídeo', 'clipe', 'cena', 'tela', 'o usuário' ou 'o mouse'.\n"
        "8. Foque nos comandos, menus e opções relevantes para reproduzir o processo."
    )
    
    base_parts = [system_instruction, time_instruction, formatting_instruction]
    
    if audio_context:
        base_parts.append(
            f"CONTEXTO DE ÁUDIO (transcrição do que está sendo dito no vídeo):\n{audio_context}"
        )
    
    if include_timestamp:
        base_parts.append(f"O procedimento ocorre especificamente ao redor de {timestamp}.")
    else:
        base_parts.append("Descreva o procedimento exibido. Não inclua timestamp na resposta.")
    
    prompt = "\n\n".join(base_parts)
    
    if user_prompt and str(user_prompt).strip():
        prompt += "\n\n" + "Contexto extra do usuário (se aplicável):\n" + str(user_prompt).strip()
    
    # 4. Call vision API
    return _analyze_frames_with_groq(frames, prompt, api_key, model_name)


def _describe_gif_with_groq(
    *,
    gif_bytes: bytes,
    api_key: str,
    model: str | None = None,
    document_context: str | None = None,
    document_title: str | None = None,
) -> str:
    """Generate a single step description from a GIF using Groq Vision.
    
    Extracts frames from the GIF, sends them to Groq with a prompt asking
    for one imperative step caption, optionally consistent with document_context.
    """
    frames = _extract_frames_from_gif(gif_bytes, frame_count=min(5, GROQ_FRAME_COUNT))
    if not frames:
        raise HTTPException(status_code=400, detail="Nenhum frame extraído do GIF")

    system_instruction = (
        f"Você é um especialista em Documentação Técnica de Software. "
        f"Seu objetivo é criar tutoriais passo a passo claros, diretos e CONTEXTUALIZADOS. "
        f"Responda sempre em português do Brasil (pt-BR). Idioma preferido: {AI_LANGUAGE}."
    )
    task_instruction = (
        "Você está vendo uma sequência de imagens extraídas de um GIF animado que mostra uma ação sendo realizada em um software. "
        "Analise CUIDADOSAMENTE a sequência para entender:\n"
        "1. Em qual tela/janela/modal o usuário está (identifique pelo título da janela, abas, cabeçalhos visíveis).\n"
        "2. Quais elementos visuais estão presentes (botões, campos de texto, menus, abas, tabelas, checkboxes).\n"
        "3. Qual é a ação específica sendo executada (clique, preenchimento, seleção, navegação).\n"
        "4. Qual é o objetivo final dessa ação."
    )
    formatting_instruction = (
        "REGRAS DE SAÍDA:\n"
        "1. Gere UMA instrução completa e contextualizada que descreva o passo/ação mostrada.\n"
        "2. INCLUA o contexto visual (nome da tela, seção, aba) no início da instrução.\n"
        "   - Exemplo RUIM: 'Clique em Salvar.'\n"
        "   - Exemplo BOM: 'Na aba \"Informações Gerais\" do cadastro de produto, clique no botão \"Salvar\" para confirmar as alterações.'\n"
        "   - Exemplo BOM: 'No campo \"NCM\", digite o código NCM do produto e pressione Enter para buscar.'\n"
        "   - Exemplo BOM: 'Acesse o menu \"Cadastros\" e selecione a opção \"Produtos\" para abrir a lista de produtos.'\n"
        "3. Use modo imperativo (ex: 'Clique em', 'Selecione', 'Preencha o campo', 'Acesse').\n"
        "4. Não use narração (evite 'O usuário clica...').\n"
        "5. Se houver textos visíveis em botões/campos/abas, mencione-os entre aspas.\n"
        "6. Não mencione 'vídeo', 'GIF' ou 'imagem' na resposta.\n"
        "7. Retorne APENAS o texto da instrução, sem numeração, prefixo ou explicações adicionais."
    )
    parts = [system_instruction, task_instruction, formatting_instruction]

    if document_context and document_context.strip():
        ctx = f"CONTEXTO DO DOCUMENTO DE REFERÊNCIA (mantenha o estilo e a estrutura):\n{document_context.strip()}"
        parts.append(ctx)
    if document_title and document_title.strip():
        parts.append(f"Título do documento: {document_title.strip()}")

    prompt = "\n\n".join(parts)
    return _analyze_frames_with_groq(frames, prompt, api_key, model)


def _is_groq_model(model_name: str) -> bool:
    """Check if the model name refers to a Groq model."""
    if not model_name:
        return False
    lower = model_name.lower()
    return (
        "llama-4" in lower
        or "llama4" in lower
        or lower.startswith("groq/")
        or lower.startswith("meta-llama/llama-4")
        or "scout" in lower
        or "maverick" in lower
    )


@dataclass
class StepPayload:
    description: str
    has_image: bool = True


def _parse_steps(steps_raw: str) -> list[StepPayload]:
    try:
        payload: Any = json.loads(steps_raw)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid steps JSON: {exc.msg}") from exc

    if not isinstance(payload, list):
        raise HTTPException(status_code=400, detail="Steps must be a JSON array")

    steps: list[StepPayload] = []
    for item in payload:
        if isinstance(item, str):
            steps.append(StepPayload(description=item, has_image=True))
            continue
        if isinstance(item, dict):
            description = item.get("description", "")
            if not isinstance(description, str):
                raise HTTPException(status_code=400, detail="Each step description must be a string")

            has_image_raw = item.get("has_image", True)
            has_image = True
            if isinstance(has_image_raw, bool):
                has_image = has_image_raw
            elif has_image_raw is None:
                has_image = True
            else:
                raise HTTPException(status_code=400, detail="Each step has_image must be a boolean")

            steps.append(StepPayload(description=description, has_image=has_image))
            continue
        raise HTTPException(status_code=400, detail="Each step must be a string or an object")

    return steps


_PNG_MAGIC = b"\x89PNG\r\n\x1a\n"
_JPEG_MAGIC = b"\xff\xd8\xff"


def _normalize_image_to_png(image_bytes: bytes) -> bytes:
    """Accept PNG as-is; accept JPEG and convert to PNG. Raise HTTPException for invalid data."""
    if not image_bytes or len(image_bytes) < 8:
        raise HTTPException(status_code=400, detail="One or more images are invalid (expected PNG or JPEG)")
    if image_bytes.startswith(_PNG_MAGIC):
        return image_bytes
    if image_bytes.startswith(_JPEG_MAGIC):
        try:
            from PIL import Image

            img = Image.open(io.BytesIO(image_bytes))
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            return buf.getvalue()
        except Exception as exc:
            raise HTTPException(
                status_code=400,
                detail="One or more images are invalid (could not convert to PNG)",
            ) from exc
    raise HTTPException(status_code=400, detail="One or more images are invalid (expected PNG or JPEG)")


def _sanitize_image_prefix(prefix: str | None) -> str:
    raw = (prefix or "").strip()
    if not raw:
        return "step_"
    # If user typed an extension, drop it to avoid "foo.png01.png"
    if raw.lower().endswith(".png"):
        raw = raw[: -len(".png")]
    # Prevent path traversal / nested paths
    raw = raw.replace("/", "_").replace("\\", "_")
    raw = "_".join(raw.split())
    allowed = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789._-")
    cleaned = "".join((ch if ch in allowed else "_") for ch in raw)
    cleaned = cleaned[:60].strip("._")
    return cleaned or "step_"




@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/videos")
async def upload_video(
    video: UploadFile = File(...),
    x_google_api_key: str | None = Header(default=None, alias="X-Google-Api-Key"),
):
    filename = (video.filename or "").strip()
    ext = Path(filename).suffix.lower()
    content_type = (video.content_type or "").lower().strip()

    content_type_to_ext = {
        "video/mp4": ".mp4",
        "video/x-matroska": ".mkv",
        "application/x-matroska": ".mkv",
    }

    if ext == "" and content_type in content_type_to_ext:
        ext = content_type_to_ext[content_type]

    if ext not in {".mp4", ".mkv"}:
        raise HTTPException(
            status_code=400,
            detail=f"Formato inválido. Use .mp4 ou .mkv (filename='{filename}', content_type='{content_type}')",
        )

    video_id = uuid.uuid4().hex
    target_path = DATA_DIR / f"video_{video_id}{ext}"

    total = 0
    try:
        with target_path.open("wb") as f:
            while True:
                chunk = await video.read(1024 * 1024)
                if not chunk:
                    break
                total += len(chunk)
                if total > MAX_VIDEO_BYTES:
                    raise HTTPException(
                        status_code=413,
                        detail=(
                            f"Vídeo muito grande (limite atual: {MAX_VIDEO_MB} MB). "
                            "Ajuste CLIPBUILDER_MAX_VIDEO_BYTES no backend/.env."
                        ),
                    )
                f.write(chunk)
    except OSError as exc:
        if getattr(exc, "errno", None) == ENOSPC:
            logger.error("disk full while saving upload to %s", target_path, exc_info=True)
            raise HTTPException(
                status_code=507,
                detail=(
                    "Sem espaço em disco para salvar o vídeo. "
                    "Ajuste CLIPBUILDER_DATA_DIR para um caminho com espaço suficiente."
                ),
            ) from exc
        logger.error("os error while saving upload to %s", target_path, exc_info=True)
        raise
    finally:
        try:
            await video.close()
        except Exception:
            pass

    # For very large videos, we keep the original locally and only send a short clip to Gemini on demand.
    with _videos_lock:
        _videos[video_id] = VideoEntry(path=target_path, status="ready")

    # Cleanup old files to stay within MAX_DATA_FILES limit
    _cleanup_old_files()

    # Validate API key early so the user gets fast feedback.
    _get_api_key(x_google_api_key)
    return {"video_id": video_id, "status": "ready"}


@app.post("/videos/raw")
async def upload_video_raw(
    request: Request,
    x_filename: str | None = Header(default=None, alias="X-Filename"),
    x_google_api_key: str | None = Header(default=None, alias="X-Google-Api-Key"),
):
    filename = (x_filename or "video.mp4").strip()
    ext = Path(filename).suffix.lower()
    if ext not in {".mp4", ".mkv"}:
        raise HTTPException(status_code=400, detail="Formato inválido. Use .mp4 ou .mkv")

    video_id = uuid.uuid4().hex
    target_path = DATA_DIR / f"video_{video_id}{ext}"

    total = 0
    try:
        with target_path.open("wb") as f:
            async for chunk in request.stream():
                if not chunk:
                    continue
                total += len(chunk)
                if total > MAX_VIDEO_BYTES:
                    raise HTTPException(
                        status_code=413,
                        detail=(
                            f"Vídeo muito grande (limite atual: {MAX_VIDEO_MB} MB). "
                            "Ajuste CLIPBUILDER_MAX_VIDEO_BYTES no backend/.env."
                        ),
                    )
                f.write(chunk)
    except HTTPException:
        raise
    except OSError as exc:
        if getattr(exc, "errno", None) == ENOSPC:
            logger.error("disk full while saving raw upload to %s", target_path, exc_info=True)
            raise HTTPException(
                status_code=507,
                detail=(
                    "Sem espaço em disco para salvar o vídeo. "
                    "Ajuste CLIPBUILDER_DATA_DIR para um caminho com espaço suficiente."
                ),
            ) from exc
        logger.error("os error while saving raw upload to %s", target_path, exc_info=True)
        raise
    except Exception as exc:
        logger.warning("raw upload stream failed: %s", exc, exc_info=True)
        raise HTTPException(status_code=400, detail="Falha ao ler o upload (conexão interrompida?)") from exc

    # For very large videos, we keep the original locally and only send a short clip to Gemini on demand.
    with _videos_lock:
        _videos[video_id] = VideoEntry(path=target_path, status="ready")

    # Cleanup old files to stay within MAX_DATA_FILES limit
    _cleanup_old_files()

    # Validate API key early so the user gets fast feedback.
    _get_api_key(x_google_api_key)
    return {"video_id": video_id, "status": "ready"}


@app.get("/videos/{video_id}/status")
def video_status(video_id: str):
    with _videos_lock:
        entry = _videos.get(video_id)
        if not entry:
            entry = _restore_video_entry_if_missing(video_id)
        if not entry:
            raise HTTPException(status_code=404, detail="Vídeo não encontrado")
        return {"status": entry.status, "error": entry.error}


@app.get("/videos/{video_id}/file")
def video_file(video_id: str):
    with _videos_lock:
        entry = _videos.get(video_id)
        if not entry:
            entry = _restore_video_entry_if_missing(video_id)
        if not entry:
            raise HTTPException(status_code=404, detail="Vídeo não encontrado")
        if entry.status != "ready":
            raise HTTPException(status_code=409, detail=entry.error or "Vídeo não está pronto")
        path = entry.path

    ext = path.suffix.lower()
    media_type = "video/mp4" if ext == ".mp4" else "video/x-matroska"
    return FileResponse(
        path=path,
        media_type=media_type,
        headers={
            "Content-Disposition": f'inline; filename="{path.name}"',
            "Cross-Origin-Resource-Policy": "cross-origin",
        },
    )


@app.post("/videos/youtube")
async def upload_youtube(
    background_tasks: BackgroundTasks,
    request: Request,
    x_google_api_key: str | None = Header(default=None, alias="X-Google-Api-Key"),
):
    # Expect JSON: {"url": "https://..."}
    try:
        payload = await request.json()
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Body inválido. Envie JSON com campo 'url'.") from exc

    url = (payload.get("url") if isinstance(payload, dict) else "")
    url = (url or "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="Campo 'url' é obrigatório")
    if not _is_supported_youtube_url(url):
        raise HTTPException(status_code=400, detail="URL não suportada. Use um link do YouTube (youtube.com / youtu.be).")

    # Validate API key early so the user gets fast feedback.
    _get_api_key(x_google_api_key)

    video_id = uuid.uuid4().hex
    target_path = DATA_DIR / f"video_{video_id}.mp4"

    with _videos_lock:
        _videos[video_id] = VideoEntry(path=target_path, status="processing")

    def job() -> None:
        try:
            _download_youtube_video(url=url, out_path=target_path)
            with _videos_lock:
                current = _videos.get(video_id)
                if current:
                    current.status = "ready"
                    current.error = None
            # Cleanup old files to stay within MAX_DATA_FILES limit
            _cleanup_old_files()
        except HTTPException as exc:
            with _videos_lock:
                current = _videos.get(video_id)
                if current:
                    current.status = "error"
                    current.error = str(exc.detail)
        except Exception as exc:
            with _videos_lock:
                current = _videos.get(video_id)
                if current:
                    current.status = "error"
                    current.error = str(exc)
            logger.error("youtube download failed (video_id=%s): %s", video_id, exc, exc_info=True)

    background_tasks.add_task(job)
    return {"video_id": video_id, "status": "processing"}


@app.get("/videos/{video_id}/smart-text")
async def smart_text(
    video_id: str,
    timestamp: str | None = None,
    t: float | None = None,
    model: str = DEFAULT_GEMINI_MODEL,
    prompt: str | None = None,
    include_timestamp: bool = True,
    x_google_api_key: str | None = Header(default=None, alias="X-Google-Api-Key"),
    x_groq_api_key: str | None = Header(default=None, alias="X-Groq-Api-Key"),
):
    with _videos_lock:
        entry = _videos.get(video_id)
        if not entry:
            entry = _restore_video_entry_if_missing(video_id)
        if not entry:
            raise HTTPException(status_code=404, detail="Vídeo não encontrado")
        if entry.status != "ready":
            raise HTTPException(status_code=409, detail=entry.error or "Vídeo não está pronto")
        source_path = entry.path

    api_key = _get_api_key(x_google_api_key)
    if timestamp is None:
        if t is None:
            raise HTTPException(status_code=400, detail="Informe timestamp (HH:MM:SS) ou t (segundos)")
        timestamp = _format_timestamp(float(t))

    try:
        ts_seconds = _parse_timestamp_to_seconds(str(timestamp))
    except ValueError:
        raise HTTPException(status_code=400, detail="timestamp inválido. Use HH:MM:SS")

    # Route to Groq if model is Llama 4 / Scout / Maverick
    if _is_groq_model(model):
        groq_key = _get_groq_api_key(x_groq_api_key)
        
        def groq_work() -> str:
            return _describe_with_groq(
                video_path=source_path,
                timestamp=str(timestamp),
                timestamp_seconds=ts_seconds,
                api_key=groq_key,
                model_name=model,
                user_prompt=prompt,
                include_timestamp=bool(include_timestamp),
            )
        
        try:
            text = await anyio.to_thread.run_sync(groq_work)
        except HTTPException:
            raise
        except Exception as exc:
            logger.error(
                "groq smart-text failed (video_id=%s, model=%s, timestamp=%s): %s",
                video_id,
                model,
                timestamp,
                exc,
                exc_info=True,
            )
            raise HTTPException(status_code=500, detail=f"Falha ao analisar com Groq: {str(exc)[:200]}") from exc
        
        return {"text": text}

    # Default: Use Gemini
    api_key = _get_api_key(x_google_api_key)
    clip_key = f"{int(ts_seconds)}:{int(GEMINI_CLIP_SECONDS)}"
    with _videos_lock:
        cached = entry.clip_cache.get(clip_key)

    if not cached:
        clip_path = DATA_DIR / f"clip_{video_id}_{int(ts_seconds)}_{int(GEMINI_CLIP_SECONDS)}.mp4"
        try:
            _make_gemini_clip(
                source_path=source_path,
                timestamp_seconds=ts_seconds,
                clip_seconds=GEMINI_CLIP_SECONDS,
                out_path=clip_path,
            )

            def upload_work() -> str:
                return _upload_video_to_gemini(clip_path, api_key)

            cached = await anyio.to_thread.run_sync(upload_work)
            with _videos_lock:
                current = _videos.get(video_id)
                if current and current.status == "ready":
                    current.clip_cache[clip_key] = cached
        except HTTPException:
            raise
        except Exception as exc:
            with _videos_lock:
                current = _videos.get(video_id)
                if current:
                    current.error = str(exc)
            logger.error(
                "gemini clip upload failed (video_id=%s, timestamp=%s): %s",
                video_id,
                timestamp,
                exc,
                exc_info=True,
            )
            raise HTTPException(status_code=500, detail="Falha ao preparar/enviar clipe para o Gemini") from exc
        finally:
            try:
                if clip_path.exists():
                    clip_path.unlink()
            except Exception:
                pass

    def work() -> str:
        return _describe_at_timestamp(
            gemini_file_name=str(cached),
            timestamp=str(timestamp),
            clip_seconds=int(GEMINI_CLIP_SECONDS),
            api_key=api_key,
            model_name=model,
            user_prompt=prompt,
            include_timestamp=bool(include_timestamp),
        )

    try:
        text = await anyio.to_thread.run_sync(work)
    except Exception as exc:
        if isinstance(exc, HTTPException):
            raise

        # Normalize common Gemini errors so the frontend gets a meaningful status.
        try:
            import re
            from google.api_core import exceptions as gexc  # type: ignore

            def _retry_hint_from_message(message: str) -> str:
                # Example: "Please retry in 13.644857575s."
                m = re.search(r"retry in\s+([0-9]+(?:\.[0-9]+)?)s", message, flags=re.IGNORECASE)
                if not m:
                    return ""
                seconds = m.group(1)
                return f" Tente novamente em ~{seconds}s."

            def _as_list(e: BaseException) -> list[BaseException]:
                # AnyIO / Python may raise ExceptionGroup with multiple nested exceptions.
                eg = getattr(e, "exceptions", None)
                if isinstance(eg, list):
                    return eg
                return [e]

            candidates: list[BaseException] = []
            for one in _as_list(exc):
                candidates.append(one)
                cause = getattr(one, "__cause__", None)
                if isinstance(cause, BaseException):
                    candidates.append(cause)

            def _msg(e: BaseException) -> str:
                try:
                    return str(e) or e.__class__.__name__
                except Exception:
                    return e.__class__.__name__

            for one in candidates:
                message = _msg(one)
                if isinstance(one, getattr(gexc, "ResourceExhausted", ())) or (
                    "quota" in message.lower() or "resourceexhausted" in message.lower() or " 429" in message
                ):
                    raise HTTPException(
                        status_code=429,
                        detail=(
                            "Quota/limite do Gemini excedido (429). Verifique billing/limites do projeto e tente novamente."
                            + _retry_hint_from_message(message)
                        ),
                    ) from exc

            if isinstance(exc, getattr(gexc, "TooManyRequests", ())):
                raise HTTPException(
                    status_code=429,
                    detail="Muitas requisições ao Gemini (429). Tente novamente em instantes.",
                ) from exc

            if isinstance(exc, getattr(gexc, "PermissionDenied", ())):
                raise HTTPException(
                    status_code=403,
                    detail="Permissão negada pelo Gemini. Verifique se a API está habilitada e se a chave é válida.",
                ) from exc

            if isinstance(exc, getattr(gexc, "NotFound", ())):
                raise HTTPException(
                    status_code=400,
                    detail="Modelo do Gemini não encontrado/suportado. Ajuste o model para um valor retornado por list_models().",
                ) from exc

            if isinstance(exc, getattr(gexc, "InvalidArgument", ())):
                raise HTTPException(status_code=400, detail="Parâmetros inválidos ao chamar o Gemini.") from exc
        except Exception:
            # If google-api-core isn't available or anything goes wrong while mapping,
            # fall back to the generic handler below.
            pass

        with _videos_lock:
            current = _videos.get(video_id)
            if current:
                current.error = str(exc)
        logger.error(
            "gemini smart-text failed (video_id=%s, model=%s, timestamp=%s): %s",
            video_id,
            model,
            timestamp,
            exc,
            exc_info=True,
        )
        raise HTTPException(status_code=500, detail="Falha ao analisar com Gemini") from exc

    return {"text": text}


@app.post("/export")
async def export_documentation(
    steps: str = Form(...),
    images: list[UploadFile] = File(default=[]),
    image_name_prefix: str | None = Form(default=None),
    output_format: str | None = Form(default="markdown"),
):
    parsed_steps = _parse_steps(steps)

    fmt = (output_format or "markdown").strip().lower()
    if fmt not in {"markdown", "html", "docx", "plain", "pdf"}:
        raise HTTPException(status_code=400, detail="Invalid output_format. Use markdown, html, docx, pdf, or plain.")

    prefix = _sanitize_image_prefix(image_name_prefix)

    expected_images = sum(1 for s in parsed_steps if s.has_image)
    if expected_images != len(images):
        raise HTTPException(
            status_code=400,
            detail=(
                f"Images count ({len(images)}) must match steps with images ({expected_images}). "
                "Send images only for steps where has_image=true."
            ),
        )

    total_bytes = 0
    processed_images: list[bytes] = []

    for upload in images:
        image_bytes = await upload.read()
        if not image_bytes:
            raise HTTPException(status_code=400, detail="One or more images are empty")

        if len(image_bytes) > MAX_SINGLE_IMAGE_BYTES:
            raise HTTPException(status_code=413, detail="One or more images are too large")

        total_bytes += len(image_bytes)
        if total_bytes > MAX_TOTAL_IMAGE_BYTES:
            raise HTTPException(status_code=413, detail="Export payload too large")

        processed_images.append(_normalize_image_to_png(image_bytes))

    md_lines: list[str] = ["# Tutorial\n"]

    step_index_to_image: dict[int, bytes] = {}
    image_cursor = 0
    for step_idx, step in enumerate(parsed_steps, start=1):
        if step.has_image:
            if image_cursor >= len(processed_images):
                raise HTTPException(status_code=400, detail="Missing image for one or more steps")
            step_index_to_image[step_idx] = processed_images[image_cursor]
            image_cursor += 1

    for i, step in enumerate(parsed_steps, start=1):
        md_lines.append(f"## Passo {i}\n")
        description = step.description.strip() or "(sem descrição)"
        md_lines.append(description + "\n")
        if step.has_image:
            md_lines.append(f"![Passo {i}](./img/{prefix}{i:02d}.png)\n")

    if fmt == "plain":
        lines: list[str] = ["Tutorial", ""]
        for i, step in enumerate(parsed_steps, start=1):
            lines.append(f"Passo {i}")
            lines.append("-" * 7)
            lines.append((step.description or "").strip() or "(sem descrição)")
            if step.has_image:
                lines.append(f"[imagem: {prefix}{i:02d}.png]")
            lines.append("")

        txt = "\n".join(lines).rstrip() + "\n"
        buf = io.BytesIO(txt.encode("utf-8"))
        headers = {"Content-Disposition": 'attachment; filename="tutorial.txt"'}
        return StreamingResponse(buf, media_type="text/plain; charset=utf-8", headers=headers)

    if fmt == "html":
        parts: list[str] = []
        parts.append("<!doctype html>")
        parts.append('<html lang="pt-br">')
        parts.append("<head>")
        parts.append('<meta charset="utf-8"/>')
        parts.append('<meta name="viewport" content="width=device-width, initial-scale=1"/>')
        parts.append("<title>Tutorial</title>")
        parts.append(
            "<style>body{font-family:system-ui,-apple-system,Segoe UI,Roboto,Ubuntu,Arial,sans-serif;max-width:900px;margin:24px auto;padding:0 16px;line-height:1.5}"  # noqa: E501
            "h1{margin:0 0 16px} h2{margin:24px 0 8px} .step{margin-bottom:18px}"  # noqa: E501
            "img{max-width:100%;height:auto;border:1px solid #ddd;border-radius:8px} pre{background:#f6f8fa;padding:12px;border-radius:8px;overflow:auto}"  # noqa: E501
            ".muted{color:#666}</style>"
        )
        parts.append("</head>")
        parts.append("<body>")
        parts.append("<h1>Tutorial</h1>")
        for i, step in enumerate(parsed_steps, start=1):
            parts.append(f"<section class=\"step\">")
            parts.append(f"<h2>Passo {i}</h2>")
            desc = (step.description or "").strip() or "(sem descrição)"
            # Basic HTML escaping
            esc = (
                desc.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            parts.append(f"<p>{esc}</p>")
            if step.has_image:
                img_bytes = step_index_to_image.get(i)
                if img_bytes:
                    b64 = base64.b64encode(img_bytes).decode("ascii")
                    parts.append(f'<img alt="Passo {i}" src="data:image/png;base64,{b64}"/>')
                else:
                    parts.append('<div class="muted">(imagem ausente)</div>')
            parts.append("</section>")
        parts.append("</body></html>")

        html_bytes = "\n".join(parts).encode("utf-8")
        buf = io.BytesIO(html_bytes)
        headers = {"Content-Disposition": 'attachment; filename="tutorial.html"'}
        return StreamingResponse(buf, media_type="text/html; charset=utf-8", headers=headers)

    if fmt == "docx":
        try:
            from docx import Document
            from docx.shared import Inches
        except Exception as exc:
            raise HTTPException(
                status_code=500,
                detail="Dependências para DOCX não instaladas. Instale python-docx e Pillow no backend.",
            ) from exc

        doc = Document()
        doc.add_heading("Tutorial", level=1)
        for i, step in enumerate(parsed_steps, start=1):
            doc.add_heading(f"Passo {i}", level=2)
            desc = (step.description or "").strip() or "(sem descrição)"
            doc.add_paragraph(desc)
            if step.has_image:
                img_bytes = step_index_to_image.get(i)
                if img_bytes:
                    stream = io.BytesIO(img_bytes)
                    try:
                        doc.add_picture(stream, width=Inches(6.5))
                    except Exception:
                        # Fallback: try without width if some builds choke on sizing
                        stream.seek(0)
                        doc.add_picture(stream)

        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        headers = {"Content-Disposition": 'attachment; filename="tutorial.docx"'}
        return StreamingResponse(
            out,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers,
        )

    if fmt == "pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import Image as RLImage
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except Exception as exc:
            raise HTTPException(
                status_code=500,
                detail="Dependências para PDF não instaladas. Instale reportlab no backend.",
            ) from exc

        font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
        try:
            pdfmetrics.registerFont(TTFont("DejaVuSans", font_path))
            base_font = "DejaVuSans"
        except Exception:
            # Fallback to a built-in font if DejaVu isn't available.
            base_font = "Helvetica"

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ClipBuilderTitle",
            parent=styles["Title"],
            fontName=base_font,
        )
        h2_style = ParagraphStyle(
            "ClipBuilderH2",
            parent=styles["Heading2"],
            fontName=base_font,
            spaceBefore=12,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "ClipBuilderBody",
            parent=styles["BodyText"],
            fontName=base_font,
            leading=14,
        )

        out = io.BytesIO()
        doc = SimpleDocTemplate(
            out,
            pagesize=A4,
            leftMargin=0.8 * inch,
            rightMargin=0.8 * inch,
            topMargin=0.8 * inch,
            bottomMargin=0.8 * inch,
            title="Tutorial",
        )

        story: list[object] = []
        story.append(Paragraph("Tutorial", title_style))
        story.append(Spacer(1, 12))

        max_width = doc.width
        for i, step in enumerate(parsed_steps, start=1):
            story.append(Paragraph(f"Passo {i}", h2_style))

            desc = (step.description or "").strip() or "(sem descrição)"
            # Basic HTML escaping for Paragraph
            esc = (
                desc.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace("\n", "<br/>")
            )
            story.append(Paragraph(esc, body_style))
            story.append(Spacer(1, 8))

            if step.has_image:
                img_bytes = step_index_to_image.get(i)
                if img_bytes:
                    # Use Pillow to preserve aspect ratio.
                    try:
                        from PIL import Image as PILImage

                        img = PILImage.open(io.BytesIO(img_bytes))
                        w_px, h_px = img.size
                    except Exception:
                        w_px, h_px = (1200, 675)

                    scale = min(1.0, float(max_width) / float(w_px or 1))
                    w = (w_px or 1) * scale
                    h = (h_px or 1) * scale

                    story.append(RLImage(io.BytesIO(img_bytes), width=w, height=h))
                    story.append(Spacer(1, 10))

        doc.build(story)
        out.seek(0)
        headers = {"Content-Disposition": 'attachment; filename="tutorial.pdf"'}
        return StreamingResponse(out, media_type="application/pdf", headers=headers)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("tutorial.md", "\n".join(md_lines).strip() + "\n")
        for step_idx, image_bytes in step_index_to_image.items():
            zf.writestr(f"img/{prefix}{step_idx:02d}.png", image_bytes)

    zip_buffer.seek(0)

    headers = {"Content-Disposition": 'attachment; filename="clipbuilder_export.zip"'}
    return StreamingResponse(zip_buffer, media_type="application/zip", headers=headers)


# ---------------------------------------------------------------------------
# Document Enhancement with Groq
# ---------------------------------------------------------------------------

@app.post("/enhance-document")
async def enhance_document(
    request: Request,
    user: CurrentAuthorizedUser,
    x_google_api_key: str | None = Header(default=None, alias="X-Google-Api-Key"),
):
    """
    Transforma passos capturados em documento profissional estruturado usando Gemini.
    
    Espera JSON no body:
    {
        "title": "Título do Documento",
        "steps": [
            {"description": "Passo 1", "timestamp": "00:01:30", "has_image": true},
            {"description": "Passo 2", "timestamp": "00:02:15", "has_image": false}
        ]
    }
    
    Retorna:
    {
        "markdown": "# Documento estruturado em Markdown..."
    }
    
    O Markdown gerado segue o formato Doc Pro:
    - Cada passo com Objetivo e Procedimento
    - Callouts (> [!NOTE], > [!TIP], > [!WARNING])
    - Checklist de Verificação Final
    - Tabela de Problemas Comuns e Soluções
    """
    from enhance import enhance_document_with_gemini
    
    # Get API key (Gemini)
    api_key = _get_api_key(x_google_api_key)
    
    # Parse request body
    try:
        body = await request.json()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"JSON inválido: {exc}") from exc
    
    title = (body.get("title") or "Documento Sem Título").strip()
    steps = body.get("steps", [])
    
    if not isinstance(steps, list):
        raise HTTPException(status_code=400, detail="Campo 'steps' deve ser uma lista")
    
    if len(steps) == 0:
        raise HTTPException(status_code=400, detail="Pelo menos um passo é necessário")
    
    # Validate steps
    validated_steps = []
    for i, step in enumerate(steps):
        if isinstance(step, str):
            validated_steps.append({"description": step, "timestamp": "", "has_image": False})
        elif isinstance(step, dict):
            validated_steps.append({
                "description": step.get("description", ""),
                "timestamp": step.get("timestamp", ""),
                "has_image": step.get("has_image", False),
            })
        else:
            raise HTTPException(status_code=400, detail=f"Passo {i+1} inválido")
    
    # Call enhancement function (Gemini)
    try:
        enhanced_markdown = enhance_document_with_gemini(
            title=title,
            steps=validated_steps,
            api_key=api_key,
            model=DEFAULT_GEMINI_MODEL,
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except Exception as exc:
        logger.error("Erro ao gerar documento com Gemini: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao gerar documento: {str(exc)[:200]}"
        ) from exc
    
    return {"markdown": enhanced_markdown}



# ---------------------------------------------------------------------------
# Markdown to DOCX (conversion only, no Groq)
# ---------------------------------------------------------------------------

@app.post("/markdown-to-docx")
async def markdown_to_docx_endpoint(request: Request):
    """
    Converte markdown em ficheiro .docx.
    Body JSON: { "markdown": "...", "title": "opcional", "professional": true/false }.
    
    Se professional=true (padrão), gera documento com formatação profissional:
    - Capa, Sumário, Títulos coloridos, Callouts, Rodapé com numeração.
    
    Se professional=false, gera documento simples.
    
    Devolve o ficheiro .docx para download.
    """
    from export_from_markdown import markdown_to_docx, markdown_to_docx_pro

    try:
        body = await request.json()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"JSON inválido: {exc}") from exc

    markdown = body.get("markdown")
    if markdown is None:
        raise HTTPException(status_code=400, detail="Campo 'markdown' é obrigatório.")
    markdown = (markdown or "").strip()
    if not markdown:
        raise HTTPException(status_code=400, detail="O markdown está vazio.")

    title = (body.get("title") or "documento_profissional").strip()
    safe_title = "".join(c if c.isalnum() or c in "._- " else "_" for c in title)[:80] or "documento_profissional"
    filename = f"{safe_title}.docx"
    
    # Default to professional mode (True)
    professional = body.get("professional", True)
    if isinstance(professional, str):
        professional = professional.lower() in ("true", "1", "yes", "sim")

    try:
        if professional:
            out_bytes = markdown_to_docx_pro(markdown, title=safe_title)
        else:
            out_bytes = markdown_to_docx(markdown)
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except Exception as exc:
        logger.error("Erro ao converter markdown para docx: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao exportar: {str(exc)[:200]}",
        ) from exc

    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        io.BytesIO(out_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


# ---------------------------------------------------------------------------
# Markdown to PDF (conversion only, no Groq)
# ---------------------------------------------------------------------------

@app.post("/markdown-to-pdf")
async def markdown_to_pdf_endpoint(request: Request):
    """
    Converte markdown em ficheiro .pdf.
    Body JSON: { "markdown": "...", "title": "opcional" }.
    Devolve o ficheiro .pdf para download.
    """
    from export_from_markdown import markdown_to_pdf

    try:
        body = await request.json()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"JSON inválido: {exc}") from exc

    markdown = body.get("markdown")
    if markdown is None:
        raise HTTPException(status_code=400, detail="Campo 'markdown' é obrigatório.")
    markdown = (markdown or "").strip()
    if not markdown:
        raise HTTPException(status_code=400, detail="O markdown está vazio.")

    title = (body.get("title") or "documento").strip()
    safe_title = "".join(c if c.isalnum() or c in "._- " else "_" for c in title)[:80] or "documento"
    filename = f"{safe_title}.pdf"

    try:
        out_bytes = markdown_to_pdf(markdown)
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except Exception as exc:
        logger.error("Erro ao converter markdown para pdf: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao exportar PDF: {str(exc)[:200]}",
        ) from exc

    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        io.BytesIO(out_bytes),
        media_type="application/pdf",
        headers=headers,
    )


# ---------------------------------------------------------------------------
# Markdown to HTML (conversion only, no Groq)
# ---------------------------------------------------------------------------

@app.post("/markdown-to-html")
async def markdown_to_html_endpoint(request: Request):
    """
    Converte markdown em ficheiro .html.
    Body JSON: { "markdown": "...", "title": "opcional" }.
    Devolve o ficheiro .html para download.
    """
    try:
        import markdown2
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="Dependência 'markdown2' não instalada. Instale no backend: pip install markdown2",
        ) from exc

    try:
        body = await request.json()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"JSON inválido: {exc}") from exc

    markdown = body.get("markdown")
    if markdown is None:
        raise HTTPException(status_code=400, detail="Campo 'markdown' é obrigatório.")
    markdown = (markdown or "").strip()

    title = (body.get("title") or "Documento").strip()
    safe_title = "".join(c if c.isalnum() or c in "._- " else "_" for c in title)[:80] or "Documento"
    filename = f"{safe_title}.html"

    def _escape_html_attr(s: str) -> str:
        return (
            s.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
        )

    try:
        html_body = markdown2.markdown(
            markdown,
            extras=["fenced-code-blocks", "tables", "break-on-newline", "nl2br"],
        )
    except Exception as exc:
        logger.error("Erro ao converter markdown para html: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Erro ao converter: {str(exc)[:200]}") from exc

    html_doc = (
        "<!doctype html>\n"
        '<html lang="pt-br">\n'
        "<head>\n"
        "  <meta charset=\"utf-8\"/>\n"
        "  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\"/>\n"
        f"  <title>{_escape_html_attr(safe_title)}</title>\n"
        "  <style>"
        "body{font-family:system-ui,-apple-system,Segoe UI,Roboto,Ubuntu,Arial,sans-serif;max-width:900px;margin:24px auto;padding:0 16px;line-height:1.5;}"
        "h1{margin:0 0 16px} h2{margin:24px 0 8px} h3{margin:18px 0 6px}"
        "img{max-width:100%;height:auto;border:1px solid #ddd;border-radius:8px}"
        "pre,code{background:#f6f8fa;padding:2px 6px;border-radius:4px;font-size:0.9em}"
        "pre{overflow:auto;padding:12px}"
        "table{border-collapse:collapse;width:100%} th,td{border:1px solid #ddd;padding:8px;text-align:left}"
        "</style>\n"
        "</head>\n"
        "<body>\n"
        f"  {html_body}\n"
        "</body>\n"
        "</html>\n"
    )

    html_bytes = html_doc.encode("utf-8")
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        io.BytesIO(html_bytes),
        media_type="text/html; charset=utf-8",
        headers=headers,
    )

@app.post("/describe-gif")
async def describe_gif(
    file: UploadFile = File(...),
    document_context: str | None = Form(default=None),
    document_title: str | None = Form(default=None),
    model: str | None = Form(default=None),
    x_groq_api_key: str | None = Header(default=None, alias="X-Groq-Api-Key"),
):
    """
    Analisa um GIF e devolve uma legenda/instrução de passo (Passo X) gerada pela IA Groq.
    
    Multipart: file (GIF obrigatório), document_context (opcional), document_title (opcional).
    Header: X-Groq-Api-Key (ou GROQ_API_KEY no .env).
    Resposta: { "description": "..." }
    """
    if not file.filename or not file.filename.lower().endswith((".gif",)):
        raise HTTPException(
            status_code=400,
            detail="Envie um arquivo GIF (extensão .gif).",
        )

    gif_bytes = await file.read()
    if len(gif_bytes) > MAX_GIF_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"GIF muito grande. Máximo: {MAX_GIF_BYTES // (1024*1024)} MB.",
        )
    if len(gif_bytes) < 100:
        raise HTTPException(status_code=400, detail="Arquivo GIF inválido ou vazio.")

    api_key = _get_groq_api_key(x_groq_api_key)

    def work() -> str:
        return _describe_gif_with_groq(
            gif_bytes=gif_bytes,
            api_key=api_key,
            model=model or DEFAULT_GROQ_VISION_MODEL,
            document_context=document_context,
            document_title=document_title,
        )

    try:
        description = await anyio.to_thread.run_sync(work)
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("describe-gif failed: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Falha ao analisar GIF com Groq: {str(exc)[:200]}",
        ) from exc

    return {"description": (description or "").strip() or "GIF animado"}


# ---------------------------------------------------------------------------
# Format document like template (Groq) and export as .docx or .pdf
# ---------------------------------------------------------------------------

@app.post("/format-and-export")
async def format_and_export(
    file: UploadFile | None = File(default=None),
    markdown: str | None = Form(default=None),
    output_format: str = Form(default="docx"),
    x_groq_api_key: str | None = Header(default=None, alias="X-Groq-Api-Key"),
):
    """
    Lê documento (upload .md/.docx/.pdf ou texto markdown), formata via Groq
    segundo o template passo a passo e devolve ficheiro .docx ou .pdf.

    Form: file (opcional), markdown (opcional), output_format (docx | pdf).
    Pelo menos um de file ou markdown é obrigatório.
    """
    from document_loader import extract_text_from_document
    from enhance import format_document_like_template
    from export_from_markdown import markdown_to_docx, markdown_to_pdf

    fmt = (output_format or "docx").strip().lower()
    if fmt not in ("docx", "pdf"):
        raise HTTPException(
            status_code=400,
            detail="output_format deve ser 'docx' ou 'pdf'.",
        )

    # Extrair texto (ficheiro ou markdown)
    try:
        raw_text = await extract_text_from_document(file=file, markdown_text=markdown)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    if not raw_text or not raw_text.strip():
        raise HTTPException(
            status_code=400,
            detail="O documento está vazio ou não foi possível extrair texto.",
        )

    # Formatar com Groq (template passo a passo)
    api_key = _get_groq_api_key(x_groq_api_key)
    try:
        formatted_md = format_document_like_template(
            raw_text=raw_text,
            api_key=api_key,
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except Exception as exc:
        logger.error("Erro ao formatar documento com Groq: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao formatar documento: {str(exc)[:200]}",
        ) from exc

    if not formatted_md or not formatted_md.strip():
        raise HTTPException(
            status_code=500,
            detail="Groq devolveu resposta vazia.",
        )

    # Exportar para docx ou pdf
    try:
        if fmt == "docx":
            out_bytes = markdown_to_docx(formatted_md)
            filename = "manual.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            out_bytes = markdown_to_pdf(formatted_md)
            filename = "manual.pdf"
            media_type = "application/pdf"
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except Exception as exc:
        logger.error("Erro ao exportar documento: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao exportar: {str(exc)[:200]}",
        ) from exc

    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        io.BytesIO(out_bytes),
        media_type=media_type,
        headers=headers,
    )


# ---------------------------------------------------------------------------
# Generate Structured Documentation (Groq JSON output)
# ---------------------------------------------------------------------------

@app.post("/generate-documentation")
async def generate_documentation(
    request: Request,
    x_groq_api_key: str | None = Header(default=None, alias="X-Groq-Api-Key"),
):
    """
    Gera documentação técnica estruturada a partir de passos capturados.
    
    Body JSON:
    {
        "title": "Título do Documento",
        "steps": [{"description": "...", "has_image": true}],
        "output_format": "json" | "markdown" | "docx"
    }
    
    Retorna:
    - JSON: Objeto estruturado com visao_geral, passos, avisos, checklist, troubleshooting
    - Markdown: Documento formatado em Markdown
    - DOCX: Arquivo .docx para download
    """
    from enhance import generate_structured_documentation
    from export_from_markdown import structured_json_to_docx

    try:
        body = await request.json()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"JSON inválido: {exc}") from exc

    title = (body.get("title") or "Documentação").strip()
    if not title:
        title = "Documentação"

    steps_raw = body.get("steps", [])
    if not steps_raw or not isinstance(steps_raw, list):
        raise HTTPException(status_code=400, detail="Campo 'steps' é obrigatório e deve ser uma lista.")

    # Validate and normalize steps
    validated_steps = []
    for i, step in enumerate(steps_raw):
        if isinstance(step, str):
            validated_steps.append({"description": step, "has_image": True})
        elif isinstance(step, dict):
            desc = step.get("description", "")
            if not isinstance(desc, str):
                raise HTTPException(status_code=400, detail=f"Passo {i+1}: 'description' deve ser string.")
            validated_steps.append({
                "description": desc,
                "has_image": step.get("has_image", True),
                "timestamp": step.get("timestamp", ""),
            })
        else:
            raise HTTPException(status_code=400, detail=f"Passo {i+1} inválido.")

    # Extract images (base64) if provided
    images_raw = body.get("images")
    if images_raw and not isinstance(images_raw, list):
         # If single string, wrap in list? No, explicit list required.
         logger.warning("Campo 'images' recebido mas não é lista.")
         images_raw = None

    output_format = (body.get("output_format") or "json").strip().lower()
    if output_format not in ("json", "markdown", "docx"):
        raise HTTPException(status_code=400, detail="output_format deve ser 'json', 'markdown' ou 'docx'.")

    # Get API key
    api_key = _get_groq_api_key(x_groq_api_key)

    # Generate structured documentation
    def work() -> dict:
        return generate_structured_documentation(
            title=title,
            steps=validated_steps,
            api_key=api_key,
            model=DEFAULT_GROQ_VISION_MODEL,
            images_b64=images_raw,
        )

    try:
        structured_doc = await anyio.to_thread.run_sync(work)
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except Exception as exc:
        logger.error("Erro ao gerar documentação estruturada: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao gerar documentação: {str(exc)[:200]}"
        ) from exc

    # Return based on output_format
    if output_format == "json":
        return {"documentation": structured_doc}

    if output_format == "markdown":
        # Convert JSON to Markdown
        md_lines = []
        md_lines.append(f"# {title}")
        md_lines.append("")
        md_lines.append("## Visão Geral")
        md_lines.append(structured_doc.get("visao_geral", ""))
        md_lines.append("")
        md_lines.append("---")
        md_lines.append("")
        
        for passo in structured_doc.get("passos", []):
            numero = passo.get("numero", 0)
            titulo = passo.get("titulo", f"Passo {numero}")
            md_lines.append(f"## Passo {numero}: {titulo}")
            md_lines.append("")
            md_lines.append("### Objetivo")
            md_lines.append(passo.get("objetivo", ""))
            md_lines.append("")
            md_lines.append("### Procedimento")
            for proc in passo.get("procedimento", []):
                md_lines.append(f"1. {proc}")
            md_lines.append("")
            if passo.get("imagem"):
                md_lines.append(f"![Passo {numero}]({passo['imagem']})")
                md_lines.append("")
            md_lines.append("---")
            md_lines.append("")
        
        # Avisos
        avisos = structured_doc.get("avisos", [])
        if avisos:
            for aviso in avisos:
                tipo = aviso.get("tipo", "NOTE").upper()
                texto = aviso.get("texto", "")
                md_lines.append(f"> [!{tipo}]")
                md_lines.append(f"> {texto}")
                md_lines.append("")
        
        # Checklist
        checklist = structured_doc.get("checklist", [])
        if checklist:
            md_lines.append("## Checklist de Verificação Final")
            md_lines.append("")
            md_lines.append("Antes de considerar o procedimento concluído, confirme:")
            md_lines.append("")
            for item in checklist:
                md_lines.append(f"- [ ] {item}")
            md_lines.append("")
            md_lines.append("---")
            md_lines.append("")
        
        # Troubleshooting
        troubleshooting = structured_doc.get("troubleshooting", [])
        if troubleshooting:
            md_lines.append("## Problemas Comuns e Soluções")
            md_lines.append("")
            md_lines.append("| Problema | Causa | Solução |")
            md_lines.append("|----------|-------|---------|")
            for item in troubleshooting:
                problema = item.get("problema", "")
                causa = item.get("causa", "")
                solucao = item.get("solucao", "")
                md_lines.append(f"| {problema} | {causa} | {solucao} |")
            md_lines.append("")
        
        return {"markdown": "\n".join(md_lines)}

    # DOCX output
    try:
        out_bytes = structured_json_to_docx(structured_doc, title=title)
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except Exception as exc:
        logger.error("Erro ao converter para DOCX: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao exportar DOCX: {str(exc)[:200]}"
        ) from exc

    safe_title = "".join(c if c.isalnum() or c in "._- " else "_" for c in title)[:80] or "documentacao"
    filename = f"{safe_title}.docx"
    
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        io.BytesIO(out_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
