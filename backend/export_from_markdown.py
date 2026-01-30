"""
Exporta Markdown (saída do Groq) para .docx ou .pdf.

Parsing simples: # ## ###, **negrito**, listas e checklist (- [ ]).
"""

from __future__ import annotations

import io
import re


def _escape_html(s: str) -> str:
    """Escapa &, <, > para uso em Paragraph do reportlab."""
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _parse_inline_bold(text: str) -> list[tuple[str, bool]]:
    """Retorna lista de (fragmento, is_bold). Ex: 'Foo **bar** baz' -> [('Foo ', False), ('bar', True), (' baz', False)]."""
    parts: list[tuple[str, bool]] = []
    rest = text
    while rest:
        i = rest.find("**")
        if i == -1:
            if rest:
                parts.append((rest, False))
            break
        before = rest[:i]
        rest = rest[i + 2:]
        j = rest.find("**")
        if j == -1:
            parts.append((before + "**" + rest, False))
            break
        if before:
            parts.append((before, False))
        parts.append((rest[:j], True))
        rest = rest[j + 2:]
    return parts


def markdown_to_docx(markdown: str) -> bytes:
    """
    Converte Markdown em bytes de um documento Word (.docx).

    Suporta: # ## ###, **negrito**, listas -, *, 1., e - [ ] (checklist).
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(
            "Dependência 'python-docx' não instalada. Instale no backend."
        ) from exc

    doc = Document()
    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # Horizontal rule / ---
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # Checklist: - [ ] ou - [x]
        if re.match(r"^[\s]*-\s*\[[\sxX]\]\s*", stripped):
            text = re.sub(r"^[\s]*-\s*\[[\sxX]\]\s*", "☐ ", stripped)
            p = doc.add_paragraph()
            for frag, bold in _parse_inline_bold(text):
                run = p.add_run(frag)
                run.bold = bold
            i += 1
            continue

        # Unordered list: - item or * item
        if re.match(r"^[\s]*[-*]\s+", stripped):
            text = re.sub(r"^[\s]*[-*]\s+", "", stripped, count=1)
            p = doc.add_paragraph(style="List Bullet")
            for frag, bold in _parse_inline_bold(text):
                run = p.add_run(frag)
                run.bold = bold
            i += 1
            continue

        # Numbered list: 1. item
        num_match = re.match(r"^[\s]*(\d+)\.\s+", stripped)
        if num_match:
            text = stripped[num_match.end() :]
            p = doc.add_paragraph(style="List Number")
            for frag, bold in _parse_inline_bold(text):
                run = p.add_run(frag)
                run.bold = bold
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        for frag, bold in _parse_inline_bold(stripped):
            run = p.add_run(frag)
            run.bold = bold
        i += 1

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def markdown_to_docx_pro(markdown: str, title: str = "Documento") -> bytes:
    """
    Converte Markdown em bytes de um documento Word (.docx) com formatação profissional.

    Recursos:
    - Layout: US Letter, margens 1" (2.54cm), fonte Arial
    - Capa centrada na primeira página
    - Sumário automático (TOC) na segunda página
    - H1: Azul Escuro + quebra de página obrigatória
    - H2/H3: Tons de azul mais claros
    - Imagens: Centralizadas, max 15cm
    - Callouts: Caixas com shading (NOTE=Azul, TIP=Verde, WARNING=Amarelo, CAUTION=Vermelho)
    - Rodapé: "Página X de Y"
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn, nsmap
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise RuntimeError(
            "Dependência 'python-docx' não instalada. Instale no backend."
        ) from exc

    doc = Document()

    # -------------------------------------------------------------------------
    # 1. Page Setup: US Letter, 1 inch margins
    # -------------------------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # -------------------------------------------------------------------------
    # 2. Set default font to Arial
    # -------------------------------------------------------------------------
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Also set Arial for headings
    for heading_level in range(1, 4):
        heading_style = doc.styles[f'Heading {heading_level}']
        heading_style.font.name = 'Arial'

    # -------------------------------------------------------------------------
    # 3. Define heading colors
    # -------------------------------------------------------------------------
    DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)      # #1F4E79 - H1
    MEDIUM_BLUE = RGBColor(0x2E, 0x75, 0xB6)    # #2E75B6 - H2
    LIGHT_BLUE = RGBColor(0x5B, 0x9B, 0xD5)     # #5B9BD5 - H3

    # Callout shading colors (XML hex format without #)
    CALLOUT_COLORS = {
        'NOTE': 'DBEAFE',      # Light blue
        'TIP': 'D1FAE5',       # Light green
        'WARNING': 'FEF3C7',   # Light yellow
        'CAUTION': 'FEE2E2',   # Light red
    }

    # -------------------------------------------------------------------------
    # Helper: Add shading to paragraph
    # -------------------------------------------------------------------------
    def _add_shading(paragraph, color_hex: str):
        """Add background shading to a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        pPr.append(shd)

    # -------------------------------------------------------------------------
    # Helper: Add page break before paragraph
    # -------------------------------------------------------------------------
    def _add_page_break_before(paragraph):
        """Force a page break before this paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        pb = OxmlElement('w:pageBreakBefore')
        pb.set(qn('w:val'), 'true')
        pPr.append(pb)

    # -------------------------------------------------------------------------
    # Helper: Create footer with Page X of Y
    # -------------------------------------------------------------------------
    def _add_page_number_footer(section):
        """Add 'Página X de Y' footer."""
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # "Página "
        run = p.add_run("Página ")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        
        # PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run2 = p.add_run()
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)
        
        # " de "
        run3 = p.add_run(" de ")
        run3.font.name = 'Arial'
        run3.font.size = Pt(10)
        
        # NUMPAGES field
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = "NUMPAGES"
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        
        run4 = p.add_run()
        run4._r.append(fldChar3)
        run4._r.append(instrText2)
        run4._r.append(fldChar4)

    # -------------------------------------------------------------------------
    # Helper: Add TOC field
    # -------------------------------------------------------------------------
    def _add_toc(doc):
        """Add Table of Contents field."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # Placeholder text
        placeholder = p.add_run("(Atualize o Sumário no Word: Ctrl+A, F9)")
        placeholder.font.italic = True
        placeholder.font.size = Pt(10)
        placeholder.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        
        run2 = p.add_run()
        run2._r.append(fldChar3)

    # -------------------------------------------------------------------------
    # 4. Add footer with page numbers
    # -------------------------------------------------------------------------
    _add_page_number_footer(section)

    # -------------------------------------------------------------------------
    # 5. Cover Page
    # -------------------------------------------------------------------------
    # Add vertical space
    for _ in range(8):
        doc.add_paragraph()
    
    # Title centered
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    
    # Subtitle/date
    doc.add_paragraph()
    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = cover_sub.add_run("Documentação Técnica")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = MEDIUM_BLUE
    
    # Page break after cover
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 6. Table of Contents Page
    # -------------------------------------------------------------------------
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_toc = toc_heading.add_run("Sumário")
    run_toc.font.name = 'Arial'
    run_toc.font.size = Pt(24)
    run_toc.font.bold = True
    run_toc.font.color.rgb = DARK_BLUE
    
    doc.add_paragraph()
    _add_toc(doc)
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 7. Parse and render Markdown content
    # -------------------------------------------------------------------------
    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    is_first_h1 = True
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            # Add a thin line or just skip
            i += 1
            continue

        # -----------------------------------------------------------------
        # Callouts: > [!NOTE], > [!TIP], > [!WARNING], > [!CAUTION]
        # -----------------------------------------------------------------
        callout_match = re.match(r"^>\s*\[!(NOTE|TIP|WARNING|CAUTION)\]", stripped, re.IGNORECASE)
        if callout_match:
            callout_type = callout_match.group(1).upper()
            callout_color = CALLOUT_COLORS.get(callout_type, 'DBEAFE')
            
            # Collect all callout lines
            callout_lines = []
            # First line might have content after [!TYPE]
            first_line_content = re.sub(r"^>\s*\[!(NOTE|TIP|WARNING|CAUTION)\]\s*", "", stripped, flags=re.IGNORECASE)
            if first_line_content:
                callout_lines.append(first_line_content)
            
            i += 1
            # Continue collecting lines that start with >
            while i < len(lines):
                next_line = lines[i].strip()
                if next_line.startswith(">"):
                    content = next_line[1:].strip()
                    if content:
                        callout_lines.append(content)
                    i += 1
                else:
                    break
            
            # Add callout box
            callout_text = " ".join(callout_lines) if callout_lines else f"{callout_type}"
            
            # Add label
            label_p = doc.add_paragraph()
            label_run = label_p.add_run(f"📌 {callout_type}: ")
            label_run.font.bold = True
            label_run.font.name = 'Arial'
            label_run.font.size = Pt(11)
            _add_shading(label_p, callout_color)
            
            # Add content
            content_p = doc.add_paragraph()
            for frag, bold in _parse_inline_bold(callout_text):
                run = content_p.add_run(frag)
                run.bold = bold
                run.font.name = 'Arial'
                run.font.size = Pt(11)
            _add_shading(content_p, callout_color)
            
            doc.add_paragraph()  # spacing after callout
            continue

        # -----------------------------------------------------------------
        # Headings
        # -----------------------------------------------------------------
        if stripped.startswith("### "):
            text = stripped[4:].strip()
            h = doc.add_heading(text, level=3)
            h.runs[0].font.color.rgb = LIGHT_BLUE
            h.runs[0].font.name = 'Arial'
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            h = doc.add_heading(text, level=2)
            h.runs[0].font.color.rgb = MEDIUM_BLUE
            h.runs[0].font.name = 'Arial'
            i += 1
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            h = doc.add_heading(text, level=1)
            h.runs[0].font.color.rgb = DARK_BLUE
            h.runs[0].font.name = 'Arial'
            # Page break before H1 (except maybe first)
            if not is_first_h1:
                _add_page_break_before(h)
            is_first_h1 = False
            i += 1
            continue

        # -----------------------------------------------------------------
        # Images: ![alt](path)
        # -----------------------------------------------------------------
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            
            # Try to add image if path is accessible
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Check if image exists (for local paths)
            from pathlib import Path
            img_file = Path(img_path)
            if img_file.exists():
                try:
                    # Max width 15cm = ~5.9 inches
                    doc.add_picture(str(img_file), width=Cm(15))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    # If image can't be added, show placeholder
                    run = p.add_run(f"[Imagem: {alt_text or img_path}]")
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            else:
                # Image path not accessible, add placeholder
                run = p.add_run(f"[Imagem: {alt_text or img_path}]")
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            
            i += 1
            continue

        # -----------------------------------------------------------------
        # Checklist: - [ ] or - [x]
        # -----------------------------------------------------------------
        if re.match(r"^[\s]*-\s*\[[\sxX]\]\s*", stripped):
            checked = re.match(r"^[\s]*-\s*\[[xX]\]", stripped)
            text = re.sub(r"^[\s]*-\s*\[[\sxX]\]\s*", "", stripped)
            symbol = "☑" if checked else "☐"
            p = doc.add_paragraph()
            for frag, bold in _parse_inline_bold(f"{symbol} {text}"):
                run = p.add_run(frag)
                run.bold = bold
                run.font.name = 'Arial'
            i += 1
            continue

        # -----------------------------------------------------------------
        # Unordered list: - item or * item
        # -----------------------------------------------------------------
        if re.match(r"^[\s]*[-*]\s+", stripped):
            text = re.sub(r"^[\s]*[-*]\s+", "", stripped, count=1)
            p = doc.add_paragraph(style="List Bullet")
            for frag, bold in _parse_inline_bold(text):
                run = p.add_run(frag)
                run.bold = bold
                run.font.name = 'Arial'
            i += 1
            continue

        # -----------------------------------------------------------------
        # Numbered list: 1. item
        # -----------------------------------------------------------------
        num_match = re.match(r"^[\s]*(\d+)\.\s+", stripped)
        if num_match:
            text = stripped[num_match.end():]
            p = doc.add_paragraph(style="List Number")
            for frag, bold in _parse_inline_bold(text):
                run = p.add_run(frag)
                run.bold = bold
                run.font.name = 'Arial'
            i += 1
            continue

        # -----------------------------------------------------------------
        # Table: | col1 | col2 |
        # -----------------------------------------------------------------
        if stripped.startswith("|") and stripped.endswith("|"):
            table_rows = []
            while i < len(lines):
                row_line = lines[i].strip()
                if row_line.startswith("|") and row_line.endswith("|"):
                    # Skip separator rows (|---|---|)
                    if not re.match(r"^\|[-:\s|]+\|$", row_line):
                        cells = [c.strip() for c in row_line.strip("|").split("|")]
                        table_rows.append(cells)
                    i += 1
                else:
                    break
            
            if table_rows:
                num_cols = max(len(row) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text
                            # Bold header row
                            if row_idx == 0:
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.font.bold = True
                                        run.font.name = 'Arial'
                
                doc.add_paragraph()  # spacing after table
            continue

        # -----------------------------------------------------------------
        # Normal paragraph
        # -----------------------------------------------------------------
        p = doc.add_paragraph()
        for frag, bold in _parse_inline_bold(stripped):
            run = p.add_run(frag)
            run.bold = bold
            run.font.name = 'Arial'
        i += 1

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def structured_json_to_docx(data: dict, title: str = "Documento") -> bytes:
    """
    Converte JSON estruturado (saída de generate_structured_documentation) em DOCX profissional.
    
    Estrutura esperada do JSON:
    {
        "visao_geral": str,
        "passos": [{"numero": int, "titulo": str, "objetivo": str, "procedimento": list[str], "imagem": str}],
        "avisos": [{"tipo": "WARNING"|"TIP"|"NOTE", "texto": str}],
        "checklist": list[str],
        "troubleshooting": [{"problema": str, "causa": str, "solucao": str}]
    }
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise RuntimeError(
            "Dependência 'python-docx' não instalada. Instale no backend."
        ) from exc

    doc = Document()

    # -------------------------------------------------------------------------
    # Page Setup: US Letter, 1 inch margins
    # -------------------------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # -------------------------------------------------------------------------
    # Set default font to Arial
    # -------------------------------------------------------------------------
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Colors
    DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
    MEDIUM_BLUE = RGBColor(0x2E, 0x75, 0xB6)
    LIGHT_BLUE = RGBColor(0x5B, 0x9B, 0xD5)

    CALLOUT_COLORS = {
        'NOTE': 'DBEAFE',
        'TIP': 'D1FAE5',
        'WARNING': 'FEF3C7',
        'CAUTION': 'FEE2E2',
    }

    def _add_shading(paragraph, color_hex: str):
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        pPr.append(shd)

    def _add_page_number_footer(section):
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run("Página ")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run2 = p.add_run()
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)
        
        run3 = p.add_run(" de ")
        run3.font.name = 'Arial'
        run3.font.size = Pt(10)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = "NUMPAGES"
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        
        run4 = p.add_run()
        run4._r.append(fldChar3)
        run4._r.append(instrText2)
        run4._r.append(fldChar4)

    # Add footer
    _add_page_number_footer(section)

    # -------------------------------------------------------------------------
    # Cover Page
    # -------------------------------------------------------------------------
    for _ in range(6):
        doc.add_paragraph()
    
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    
    doc.add_paragraph()
    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = cover_sub.add_run("Documentação Técnica")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = MEDIUM_BLUE
    
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # Visão Geral
    # -------------------------------------------------------------------------
    h = doc.add_heading("Visão Geral", level=1)
    h.runs[0].font.color.rgb = DARK_BLUE
    h.runs[0].font.name = 'Arial'
    
    visao_geral = data.get("visao_geral", "")
    if visao_geral:
        p = doc.add_paragraph(visao_geral)
        p.runs[0].font.name = 'Arial' if p.runs else None
    
    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # Passos
    # -------------------------------------------------------------------------
    passos = data.get("passos", [])
    for passo in passos:
        numero = passo.get("numero", 0)
        titulo = passo.get("titulo", f"Passo {numero}")
        objetivo = passo.get("objetivo", "")
        procedimento = passo.get("procedimento", [])
        imagem = passo.get("imagem")
        
        # Título do passo
        h = doc.add_heading(f"Passo {numero}: {titulo}", level=2)
        h.runs[0].font.color.rgb = MEDIUM_BLUE
        h.runs[0].font.name = 'Arial'
        
        # Objetivo
        obj_label = doc.add_paragraph()
        run_label = obj_label.add_run("Objetivo")
        run_label.font.bold = True
        run_label.font.name = 'Arial'
        run_label.font.size = Pt(11)
        
        if objetivo:
            obj_p = doc.add_paragraph(objetivo)
            for r in obj_p.runs:
                r.font.name = 'Arial'
        
        # Procedimento
        proc_label = doc.add_paragraph()
        run_proc = proc_label.add_run("Procedimento")
        run_proc.font.bold = True
        run_proc.font.name = 'Arial'
        run_proc.font.size = Pt(11)
        
        if isinstance(procedimento, list):
            for step_text in procedimento:
                p = doc.add_paragraph(style="List Number")
                p.add_run(step_text).font.name = 'Arial'
        elif procedimento:
            p = doc.add_paragraph(str(procedimento))
            for r in p.runs:
                r.font.name = 'Arial'
        
        # Imagem placeholder
        if imagem:
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = img_p.add_run(f"[Imagem: {imagem}]")
            run_img.font.italic = True
            run_img.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        
        doc.add_paragraph()

    # -------------------------------------------------------------------------
    # Avisos
    # -------------------------------------------------------------------------
    avisos = data.get("avisos", [])
    if avisos:
        h = doc.add_heading("Avisos Importantes", level=2)
        h.runs[0].font.color.rgb = MEDIUM_BLUE
        h.runs[0].font.name = 'Arial'
        
        for aviso in avisos:
            tipo = aviso.get("tipo", "NOTE").upper()
            texto = aviso.get("texto", "")
            color_hex = CALLOUT_COLORS.get(tipo, 'DBEAFE')
            
            emoji = {"WARNING": "⚠️", "TIP": "💡", "CAUTION": "🚨", "NOTE": "📝"}.get(tipo, "📌")
            
            p = doc.add_paragraph()
            run = p.add_run(f"{emoji} {tipo}: {texto}")
            run.font.name = 'Arial'
            _add_shading(p, color_hex)
        
        doc.add_paragraph()

    # -------------------------------------------------------------------------
    # Checklist de Verificação Final
    # -------------------------------------------------------------------------
    checklist = data.get("checklist", [])
    if checklist:
        h = doc.add_heading("Checklist de Verificação Final", level=2)
        h.runs[0].font.color.rgb = MEDIUM_BLUE
        h.runs[0].font.name = 'Arial'
        
        intro = doc.add_paragraph("Antes de considerar o procedimento concluído, confirme:")
        for r in intro.runs:
            r.font.name = 'Arial'
        
        for item in checklist:
            p = doc.add_paragraph()
            run = p.add_run(f"☐ {item}")
            run.font.name = 'Arial'
        
        doc.add_paragraph()

    # -------------------------------------------------------------------------
    # Tabela de Troubleshooting
    # -------------------------------------------------------------------------
    troubleshooting = data.get("troubleshooting", [])
    if troubleshooting:
        h = doc.add_heading("Problemas Comuns e Soluções", level=2)
        h.runs[0].font.color.rgb = MEDIUM_BLUE
        h.runs[0].font.name = 'Arial'
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ["Problema", "Causa", "Solução"]
        for i, header_text in enumerate(headers):
            header_cells[i].text = header_text
            for p in header_cells[i].paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.name = 'Arial'
        
        # Data rows
        for item in troubleshooting:
            row = table.add_row().cells
            row[0].text = item.get("problema", "")
            row[1].text = item.get("causa", "")
            row[2].text = item.get("solucao", "")
            
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Arial'

    # Save to bytes
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def markdown_to_pdf(markdown: str) -> bytes:
    """
    Converte Markdown em bytes de um PDF.

    Reutiliza padrão do projeto: DejaVu/Helvetica, SimpleDocTemplate, Paragraph.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError(
            "Dependência 'reportlab' não instalada. Instale no backend."
        ) from exc

    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    try:
        pdfmetrics.registerFont(TTFont("DejaVuSans", font_path))
        base_font = "DejaVuSans"
    except Exception:
        base_font = "Helvetica"

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "MdTitle",
        parent=styles["Title"],
        fontName=base_font,
    )
    h2_style = ParagraphStyle(
        "MdH2",
        parent=styles["Heading2"],
        fontName=base_font,
        spaceBefore=12,
        spaceAfter=6,
    )
    h3_style = ParagraphStyle(
        "MdH3",
        parent=styles["Heading3"],
        fontName=base_font,
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "MdBody",
        parent=styles["BodyText"],
        fontName=base_font,
        leading=14,
    )

    out = io.BytesIO()
    doc = SimpleDocTemplate(
        out,
        pagesize=A4,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="Manual",
    )

    story: list[object] = []
    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("### "):
            text = _escape_html(stripped[4:].strip())
            story.append(Paragraph(text, h3_style))
            story.append(Spacer(1, 4))
            i += 1
            continue
        if stripped.startswith("## "):
            text = _escape_html(stripped[3:].strip())
            story.append(Paragraph(text, h2_style))
            story.append(Spacer(1, 6))
            i += 1
            continue
        if stripped.startswith("# "):
            text = _escape_html(stripped[2:].strip())
            story.append(Paragraph(text, title_style))
            story.append(Spacer(1, 12))
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # Checklist or list
        if re.match(r"^[\s]*-\s*\[[\sxX]\]\s*", stripped):
            text = re.sub(r"^[\s]*-\s*\[[\sxX]\]\s*", "☐ ", stripped)
            text = _escape_html(text).replace("\n", "<br/>")
            story.append(Paragraph(text, body_style))
            i += 1
            continue
        if re.match(r"^[\s]*[-*]\s+", stripped):
            text = re.sub(r"^[\s]*[-*]\s+", "• ", stripped, count=1)
            text = _escape_html(text).replace("\n", "<br/>")
            story.append(Paragraph(text, body_style))
            i += 1
            continue
        num_match = re.match(r"^[\s]*(\d+)\.\s+", stripped)
        if num_match:
            text = stripped[num_match.end() :]
            text = _escape_html(text).replace("\n", "<br/>")
            story.append(Paragraph(text, body_style))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        text = _escape_html(stripped).replace("\n", "<br/>")
        story.append(Paragraph(text, body_style))
        story.append(Spacer(1, 6))
        i += 1

    doc.build(story)
    out.seek(0)
    return out.read()
